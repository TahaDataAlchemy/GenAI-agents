import streamlit as st
import requests
from docx import Document

# FastAPI endpoints
TEXT_PROCESS_API = "http://127.0.0.1:8000/process_text"
DOWNLOAD_API = "http://127.0.0.1:8000/download_report"

st.title("ChatScribe 🤖")

# User input field
user_input = st.text_area("Enter your text:")

if st.button("Analyze Text"):
    if user_input:
        with st.spinner("Processing..."):
            response = requests.post(TEXT_PROCESS_API, json={"text": user_input})
            if response.status_code == 200:
                data = response.json()
                
                # Display results
                st.subheader("Analysis Results:")
                st.write(f"**Classification:** {data.get('classification', 'N/A')}")
                st.write(f"**Entities:** {', '.join(data.get('entities', []))}")
                st.write(f"**Summary:** {data.get('summary', 'N/A')}")

                # Create DOCX report
                doc = Document()
                doc.add_heading("Text Analysis Report", level=1)
                doc.add_paragraph(f"**Classification:** {data.get('classification', 'N/A')}")
                doc.add_paragraph(f"**Entities:** {', '.join(data.get('entities', []))}")
                doc.add_paragraph(f"**Summary:** {data.get('summary', 'N/A')}")
                
                # Save file
                report_path = "analysis_report.docx"
                doc.save(report_path)

                # Download Button
                with open(report_path, "rb") as file:
                    st.download_button("Download Report 📄", file, file_name="Text_Analysis_Report.docx")

            else:
                st.error(f"Error: {response.status_code} - {response.text}")
    else:
        st.warning("Please enter some text!")

